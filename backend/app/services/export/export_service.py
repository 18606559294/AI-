"""
导出服务 - PDF/Word/图片导出
"""
import io
import os
from typing import Dict, Any, Optional
from datetime import datetime
from jinja2 import Template as JinjaTemplate
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# PDF 导出需要 WeasyPrint
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False


class ExportService:
    """导出服务类"""

    # 可用模板样式
    AVAILABLE_STYLES = {
        "modern": {
            "name": "现代简约",
            "primary_color": "#2563EB",
            "secondary_color": "#64748B",
            "font_family": "Microsoft YaHei, sans-serif",
            "layout": "compact"
        },
        "professional": {
            "name": "商务专业",
            "primary_color": "#1E3A5F",
            "secondary_color": "#475569",
            "font_family": "SimSun, serif",
            "layout": "traditional"
        },
        "creative": {
            "name": "创意设计",
            "primary_color": "#7C3AED",
            "secondary_color": "#A78BFA",
            "font_family": "Microsoft YaHei, sans-serif",
            "layout": "spacious"
        },
        "minimal": {
            "name": "极简主义",
            "primary_color": "#000000",
            "secondary_color": "#666666",
            "font_family": "Helvetica, Arial, sans-serif",
            "layout": "minimal"
        },
        "executive": {
            "name": "高管风格",
            "primary_color": "#1a1a2e",
            "secondary_color": "#4a4a6a",
            "font_family": "Georgia, serif",
            "layout": "executive"
        },
        "tech": {
            "name": "技术蓝调",
            "primary_color": "#0EA5E9",
            "secondary_color": "#6366F1",
            "font_family": "Consolas, Monaco, monospace",
            "layout": "tech"
        },
        "academic": {
            "name": "学术风格",
            "primary_color": "#1B4D3E",
            "secondary_color": "#2E8B57",
            "font_family": "Times New Roman, serif",
            "layout": "academic"
        },
        "startup": {
            "name": "创业活力",
            "primary_color": "#FF6B35",
            "secondary_color": "#004E89",
            "font_family": "Segoe UI, sans-serif",
            "layout": "startup"
        },
        "elegant": {
            "name": "优雅典雅",
            "primary_color": "#8B7355",
            "secondary_color": "#A0826D",
            "font_family": "Palatino Linotype, serif",
            "layout": "elegant"
        },
        "fresh": {
            "name": "清新自然",
            "primary_color": "#10B981",
            "secondary_color": "#34D399",
            "font_family": "Calibri, sans-serif",
            "layout": "fresh"
        }
    }

    def __init__(self):
        self.template_dir = os.path.join(os.path.dirname(__file__), "templates")

    def get_available_styles(self) -> Dict[str, Dict[str, str]]:
        """获取可用的导出样式"""
        return self.AVAILABLE_STYLES
    
    async def to_pdf(
        self,
        resume_content: Dict[str, Any],
        style_config: Optional[Dict[str, Any]] = None,
        template_html: Optional[str] = None
    ) -> bytes:
        """导出为PDF"""
        if not WEASYPRINT_AVAILABLE:
            raise RuntimeError("PDF导出功能需要安装 WeasyPrint")

        # 生成HTML（内联CSS样式）
        html_content = self._generate_html_with_inline_css(resume_content, style_config, template_html)

        # 配置中文字体支持
        css = CSS(string=self._get_pdf_css(style_config))

        # 转换为PDF（使用WeasyPrint 60.2的正确API）
        html = HTML(string=html_content)
        pdf_bytes = html.write_pdf(stylesheets=[css])

        return pdf_bytes

    def _get_pdf_css(self, style_config: Optional[Dict[str, Any]] = None) -> str:
        """获取PDF专用CSS，包含中文字体配置"""
        config = style_config or {}
        font_family = config.get("font_family", "Microsoft YaHei, SimHei, sans-serif")

        # 为PDF添加中文字体支持和页面设置
        return f"""
        @page {{
            size: A4;
            margin: 2cm;
            @bottom-right {{
                content: "第 " counter(page) " 页";
                font-size: 9pt;
                color: #666;
            }}
        }}
        body {{
            font-family: {font_family};
        }}
        """

    def _generate_html_with_inline_css(
        self,
        resume_content: Dict[str, Any],
        style_config: Optional[Dict[str, Any]] = None,
        template_html: Optional[str] = None
    ) -> str:
        """生成内联CSS的HTML内容"""
        if template_html:
            template = JinjaTemplate(template_html)
            return template.render(resume=resume_content, style=style_config or {})

        # 生成HTML和CSS
        html_body = self._generate_html(resume_content, style_config, template_html)
        css_content = self._generate_css(style_config)

        # 组合完整HTML
        full_html = f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{resume_content.get('basic_info', {}).get('name', '简历')}</title>
    <style>
    {css_content}
    </style>
</head>
<body>
{html_body}
</body>
</html>
        """
        return full_html
    
    async def to_word(
        self,
        resume_content: Dict[str, Any],
        style_config: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """导出为Word文档"""
        doc = Document()
        
        # 设置文档样式
        self._setup_word_styles(doc, style_config)
        
        # 基本信息
        basic_info = resume_content.get("basic_info", {})
        if basic_info:
            self._add_basic_info_section(doc, basic_info)
        
        # 个人简介
        if basic_info.get("self_introduction"):
            self._add_section(doc, "个人简介", basic_info["self_introduction"])
        
        # 教育经历
        education = resume_content.get("education", [])
        if education:
            self._add_education_section(doc, education)
        
        # 工作经历
        work_experience = resume_content.get("work_experience", [])
        if work_experience:
            self._add_work_section(doc, work_experience)
        
        # 项目经历
        projects = resume_content.get("projects", [])
        if projects:
            self._add_project_section(doc, projects)
        
        # 技能
        skills = resume_content.get("skills", [])
        if skills:
            self._add_skills_section(doc, skills)
        
        # 证书
        certifications = resume_content.get("certifications", [])
        if certifications:
            self._add_certifications_section(doc, certifications)
        
        # 保存到字节流
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.read()
    
    async def to_html(
        self,
        resume_content: Dict[str, Any],
        style_config: Optional[Dict[str, Any]] = None,
        template_html: Optional[str] = None
    ) -> str:
        """导出为HTML"""
        html_content = self._generate_html(resume_content, style_config, template_html)
        css_content = self._generate_css(style_config)
        
        # 组合完整HTML
        full_html = f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{resume_content.get('basic_info', {}).get('name', '简历')}</title>
    <style>
    {css_content}
    </style>
</head>
<body>
{html_content}
</body>
</html>
        """
        return full_html
    
    def _generate_html(
        self,
        resume_content: Dict[str, Any],
        style_config: Optional[Dict[str, Any]] = None,
        template_html: Optional[str] = None
    ) -> str:
        """生成HTML内容"""
        if template_html:
            template = JinjaTemplate(template_html)
            return template.render(resume=resume_content, style=style_config or {})
        
        # 默认HTML模板
        basic_info = resume_content.get("basic_info", {})
        
        html_parts = ['<div class="resume">']
        
        # 头部信息
        html_parts.append('<header class="resume-header">')
        if basic_info.get("name"):
            html_parts.append(f'<h1 class="name">{basic_info["name"]}</h1>')
        if basic_info.get("job_intention"):
            html_parts.append(f'<p class="job-intention">{basic_info["job_intention"]}</p>')
        
        # 联系方式
        contact_info = []
        if basic_info.get("phone"):
            contact_info.append(f'<span>📱 {basic_info["phone"]}</span>')
        if basic_info.get("email"):
            contact_info.append(f'<span>📧 {basic_info["email"]}</span>')
        if basic_info.get("location"):
            contact_info.append(f'<span>📍 {basic_info["location"]}</span>')
        if contact_info:
            html_parts.append(f'<div class="contact-info">{" | ".join(contact_info)}</div>')
        html_parts.append('</header>')
        
        # 个人简介
        if basic_info.get("self_introduction"):
            html_parts.append('<section class="section">')
            html_parts.append('<h2 class="section-title">个人简介</h2>')
            html_parts.append(f'<p>{basic_info["self_introduction"]}</p>')
            html_parts.append('</section>')
        
        # 教育经历
        education = resume_content.get("education", [])
        if education:
            html_parts.append('<section class="section">')
            html_parts.append('<h2 class="section-title">教育经历</h2>')
            for edu in education:
                html_parts.append('<div class="item">')
                html_parts.append(f'<div class="item-header">')
                html_parts.append(f'<strong>{edu.get("school", "")}</strong>')
                html_parts.append(f'<span>{edu.get("major", "")} | {edu.get("degree", "")}</span>')
                html_parts.append('</div>')
                if edu.get("description"):
                    html_parts.append(f'<p>{edu["description"]}</p>')
                html_parts.append('</div>')
            html_parts.append('</section>')
        
        # 工作经历
        work_experience = resume_content.get("work_experience", [])
        if work_experience:
            html_parts.append('<section class="section">')
            html_parts.append('<h2 class="section-title">工作经历</h2>')
            for work in work_experience:
                html_parts.append('<div class="item">')
                html_parts.append(f'<div class="item-header">')
                html_parts.append(f'<strong>{work.get("company", "")}</strong>')
                html_parts.append(f'<span>{work.get("position", "")}</span>')
                html_parts.append('</div>')
                if work.get("description"):
                    html_parts.append(f'<p>{work["description"]}</p>')
                achievements = work.get("achievements", [])
                if achievements:
                    html_parts.append('<ul>')
                    for achievement in achievements:
                        html_parts.append(f'<li>{achievement}</li>')
                    html_parts.append('</ul>')
                html_parts.append('</div>')
            html_parts.append('</section>')
        
        # 项目经历
        projects = resume_content.get("projects", [])
        if projects:
            html_parts.append('<section class="section">')
            html_parts.append('<h2 class="section-title">项目经历</h2>')
            for project in projects:
                html_parts.append('<div class="item">')
                html_parts.append(f'<div class="item-header">')
                html_parts.append(f'<strong>{project.get("name", "")}</strong>')
                if project.get("role"):
                    html_parts.append(f'<span>{project["role"]}</span>')
                html_parts.append('</div>')
                if project.get("description"):
                    html_parts.append(f'<p>{project["description"]}</p>')
                tech_stack = project.get("tech_stack", [])
                if tech_stack:
                    html_parts.append(f'<p class="tech-stack">技术栈: {", ".join(tech_stack)}</p>')
                html_parts.append('</div>')
            html_parts.append('</section>')
        
        # 技能
        skills = resume_content.get("skills", [])
        if skills:
            html_parts.append('<section class="section">')
            html_parts.append('<h2 class="section-title">专业技能</h2>')
            html_parts.append('<div class="skills">')
            for skill in skills:
                html_parts.append(f'<span class="skill-tag">{skill.get("name", "")}</span>')
            html_parts.append('</div>')
            html_parts.append('</section>')
        
        html_parts.append('</div>')
        
        return '\n'.join(html_parts)
    
    def _generate_css(self, style_config: Optional[Dict[str, Any]] = None) -> str:
        """生成CSS样式 - 支持预设风格和布局"""
        config = style_config or {}
        style_theme = config.get("theme", "modern")

        # 根据主题选择预设样式
        if style_theme in self.AVAILABLE_STYLES:
            theme_config = self.AVAILABLE_STYLES[style_theme]
            primary_color = theme_config["primary_color"]
            secondary_color = theme_config["secondary_color"]
            font_family = theme_config["font_family"]
            layout = theme_config.get("layout", "compact")
        else:
            primary_color = config.get("primary_color", "#2563EB")
            secondary_color = config.get("secondary_color", "#64748B")
            font_family = config.get("font_family", "Microsoft YaHei, sans-serif")
            layout = config.get("layout", "compact")

        font_size = config.get("font_size", 12)
        line_height = config.get("line_height", 1.6)

        # 根据布局类型生成不同的CSS
        layout_css = self._get_layout_css(layout, primary_color, secondary_color)
        
        return f"""
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        body {{
            font-family: {font_family};
            font-size: {font_size}pt;
            line-height: {line_height};
            color: {primary_color};
            background: #fff;
        }}
        {layout_css}
        .resume {{
            max-width: 800px;
            margin: 0 auto;
            padding: 40px;
        }}
        .resume-header {{
            text-align: center;
            margin-bottom: 30px;
            padding-bottom: 20px;
            border-bottom: 2px solid {primary_color};
        }}
        .name {{
            font-size: 28pt;
            font-weight: bold;
            margin-bottom: 5px;
        }}
        .job-intention {{
            font-size: 14pt;
            color: {secondary_color};
            margin-bottom: 10px;
        }}
        .contact-info {{
            font-size: 10pt;
            color: {secondary_color};
        }}
        .contact-info span {{
            margin: 0 10px;
        }}
        .section {{
            margin-bottom: 25px;
        }}
        .section-title {{
            font-size: 14pt;
            font-weight: bold;
            color: {primary_color};
            margin-bottom: 15px;
            padding-bottom: 5px;
            border-bottom: 1px solid #ddd;
        }}
        .item {{
            margin-bottom: 15px;
        }}
        .item-header {{
            display: flex;
            justify-content: space-between;
            margin-bottom: 5px;
        }}
        .item-header strong {{
            color: {primary_color};
        }}
        .item-header span {{
            color: {secondary_color};
        }}
        .item p {{
            color: {secondary_color};
            margin-bottom: 5px;
        }}
        .item ul {{
            margin-left: 20px;
            color: {secondary_color};
        }}
        .item li {{
            margin-bottom: 3px;
        }}
        .skills {{
            display: flex;
            flex-wrap: wrap;
            gap: 10px;
        }}
        .skill-tag {{
            background: #f5f5f5;
            padding: 5px 15px;
            border-radius: 15px;
            font-size: 10pt;
        }}
        .tech-stack {{
            font-size: 10pt;
            color: {secondary_color};
            font-style: italic;
        }}
        @media print {{
            .resume {{
                padding: 20px;
            }}
            .contact-info span {{
                margin: 0 5px;
            }}
        }}

        /* 响应式设计 - 平板 */
        @media screen and (max-width: 768px) {{
            .resume {{
                padding: 30px 20px;
                max-width: 100%;
            }}
            .name {{
                font-size: 24pt;
            }}
            .job-intention {{
                font-size: 12pt;
            }}
            .item-header {{
                flex-direction: column;
            }}
            .contact-info {{
                flex-wrap: wrap;
                justify-content: center;
            }}
            .contact-info span {{
                margin: 5px 10px;
            }}
        }}

        /* 响应式设计 - 手机 */
        @media screen and (max-width: 480px) {{
            .resume {{
                padding: 20px 15px;
            }}
            .name {{
                font-size: 20pt;
            }}
            .job-intention {{
                font-size: 11pt;
            }}
            .section-title {{
                font-size: 12pt;
            }}
            .skills {{
                flex-direction: column;
            }}
            .skill-tag {{
                display: inline-block;
                margin: 3px 0;
            }}
            .contact-info {{
                flex-direction: column;
                align-items: center;
            }}
            .contact-info span {{
                display: block;
                margin: 3px 0;
            }}
            .item-header {{
                flex-direction: column;
                align-items: flex-start;
            }}
            .item-header span {{
                margin-left: 0;
                font-size: 9pt;
            }}
        }}

        /* 打印优化 */
        @media print {{
            body {{
                background: white;
            }}
            .resume {{
                box-shadow: none;
                margin: 0;
                max-width: 100%;
            }}
            a {{
                text-decoration: none;
                color: inherit;
            }}
            .section {{
                page-break-inside: avoid;
            }}
            .item {{
                page-break-inside: avoid;
            }}
        }}
        """

    def _get_layout_css(self, layout: str, primary_color: str, secondary_color: str) -> str:
        """根据布局类型生成特定CSS"""
        layout_styles = {
            "compact": """
                .resume { max-width: 750px; }
                .section { margin-bottom: 20px; }
                .item { margin-bottom: 12px; }
            """,
            "traditional": """
                .resume { max-width: 850px; }
                .section-title { border-bottom: 2px solid %s; }
                .section { margin-bottom: 30px; }
            """ % primary_color,
            "spacious": """
                .resume { max-width: 900px; padding: 50px; }
                .section { margin-bottom: 35px; }
                .item { margin-bottom: 20px; padding: 15px 0; }
                .skill-tag { padding: 8px 20px; margin: 8px; }
            """,
            "minimal": """
                .resume { max-width: 700px; }
                .section-title { border-bottom: 1px solid #000; }
                .resume-header { border-bottom: 1px solid #000; }
            """,
            "executive": """
                .resume { max-width: 800px; }
                .resume-header { background: #f8f8f8; padding: 20px; margin: -40px -40px 30px; }
                .section-title { background: #f0f0f0; padding: 8px 15px; border: none; }
                .item { border-left: 3px solid %s; padding-left: 15px; }
            """ % primary_color,
            "tech": """
                .resume { font-family: 'Consolas', 'Monaco', monospace; }
                .section-title { font-family: 'Consolas', monospace; border-left: 4px solid %s; padding-left: 10px; }
                .skill-tag { font-family: 'Consolas', monospace; background: #1a1a1a; color: #fff; }
                .item { background: #f5f5f5; padding: 10px; border-radius: 4px; }
            """ % primary_color,
            "academic": """
                .resume { max-width: 800px; text-align: justify; }
                .section-title { text-align: center; border-bottom: 2px double %s; }
                .item-header { flex-direction: column; }
                .item-header span { font-style: italic; }
            """ % primary_color,
            "startup": """
                .resume { max-width: 800px; }
                .resume-header { border-bottom: 4px solid %s; }
                .section-title { color: #fff; background: %s; padding: 10px 15px; }
                .skill-tag { background: %s; color: #fff; }
            """ % (primary_color, primary_color, primary_color),
            "elegant": """
                .resume { max-width: 800px; }
                .resume-header { border-bottom: 3px double %s; }
                .section-title { font-family: 'Palatino Linotype', serif; letter-spacing: 1px; }
                .item-header strong { font-family: 'Palatino Linotype', serif; }
            """ % primary_color,
            "fresh": """
                .resume { max-width: 800px; }
                .resume-header { border-bottom: 3px solid %s; border-radius: 0 0 5px 5px; }
                .skill-tag { background: #e6f7f0; color: #10B981; border: 1px solid #10B981; }
                .section-title { color: %s; }
            """ % (primary_color, primary_color)
        }

        return layout_styles.get(layout, "")

    def _setup_word_styles(self, doc: Document, style_config: Optional[Dict[str, Any]] = None):
        """设置Word文档样式"""
        config = style_config or {}

        # 设置默认样式
        style = doc.styles['Normal']
        font = style.font
        font.name = config.get("font_family", "微软雅黑")
        font.size = Pt(config.get("font_size", 11))

        # 设置中文字体（兼容性）
        r = style.element
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '微软雅黑')
        rPr.append(rFonts)

        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
    
    def _add_basic_info_section(self, doc: Document, basic_info: Dict[str, Any]):
        """添加基本信息部分"""
        # 姓名
        if basic_info.get("name"):
            p = doc.add_paragraph()
            run = p.add_run(basic_info["name"])
            run.bold = True
            run.font.size = Pt(22)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 求职意向
        if basic_info.get("job_intention"):
            p = doc.add_paragraph()
            p.add_run(basic_info["job_intention"])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 联系方式
        contact_parts = []
        if basic_info.get("phone"):
            contact_parts.append(basic_info["phone"])
        if basic_info.get("email"):
            contact_parts.append(basic_info["email"])
        if basic_info.get("location"):
            contact_parts.append(basic_info["location"])
        
        if contact_parts:
            p = doc.add_paragraph()
            p.add_run(" | ".join(contact_parts))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # 空行
    
    def _add_section(self, doc: Document, title: str, content: str):
        """添加通用段落"""
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        
        doc.add_paragraph(content)
        doc.add_paragraph()
    
    def _add_education_section(self, doc: Document, education: list):
        """添加教育经历 - 改进版使用更好的格式"""
        p = doc.add_paragraph()
        run = p.add_run("教育经历")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        p.space_after = Pt(12)

        for edu in education:
            # 学校和专业
            p = doc.add_paragraph()
            school_run = p.add_run(f"{edu.get('school', '')}")
            school_run.bold = True
            school_run.font.size = Pt(11)

            if edu.get("major") or edu.get("degree"):
                major_text = f" — {edu.get('major', '')}"
                if edu.get("degree"):
                    major_text += f" ({edu['degree']})"
                p.add_run(major_text)

            # 时间段
            if edu.get("start_date") or edu.get("end_date"):
                start = edu.get("start_date", "")
                end = edu.get("end_date", "")
                p = doc.add_paragraph(f"{start} ~ {end}")
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                p.space_after = Pt(6)

            # GPA/成绩
            if edu.get("gpa"):
                p = doc.add_paragraph(f"GPA: {edu['gpa']}")
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.italic = True

            # 课程/描述
            if edu.get("description"):
                p = doc.add_paragraph(edu["description"])
                p.runs[0].font.size = Pt(10)

            if edu.get("courses"):
                courses = edu.get("courses", [])
                if isinstance(courses, list):
                    p = doc.add_paragraph()
                    p.add_run("主要课程: ").bold = True
                    p.add_run(", ".join(courses[:8]))  # 最多显示8门课程
                    p.runs[1].font.size = Pt(9)

            doc.add_paragraph()  # 空行分隔各教育经历
    
    def _add_work_section(self, doc: Document, work_experience: list):
        """添加工作经历 - 改进版使用更好的格式和分页控制"""
        p = doc.add_paragraph()
        run = p.add_run("工作经历")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        p.space_after = Pt(12)

        for idx, work in enumerate(work_experience):
            # 防止在单个工作经历中间分页
            if idx > 0:
                # 添加分页控制（保持段落不被分割）
                p = doc.add_paragraph()
                p.paragraph_format.widow_control = True
                p.paragraph_format.keep_together = True

            # 公司和职位 - 使用两行布局更清晰
            p = doc.add_paragraph()
            company_run = p.add_run(work.get('company', ''))
            company_run.bold = True
            company_run.font.size = Pt(11)
            company_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

            # 职位和时间
            p = doc.add_paragraph()
            position_run = p.add_run(work.get('position', ''))
            position_run.bold = True
            position_run.font.size = Pt(10)
            position_run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

            # 添加时间
            if work.get("start_date") or work.get("end_date"):
                start = work.get("start_date", "")
                end = work.get("end_date", "至今")
                time_run = p.add_run(f"  |  {start} ~ {end}")
                time_run.font.size = Pt(9)
                time_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

            # 工作地点
            if work.get("location"):
                p = doc.add_paragraph(f"📍 {work['location']}")
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

            # 工作描述
            if work.get("description"):
                p = doc.add_paragraph(work["description"])
                p.runs[0].font.size = Pt(10)
                p.space_after = Pt(6)

            # 主要成果 - 使用列表格式
            achievements = work.get("achievements", [])
            if achievements:
                p = doc.add_paragraph()
                run = p.add_run("主要成果:")
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

                for achievement in achievements:
                    p = doc.add_paragraph(f"• {achievement}", style="List Bullet")
                    p.runs[0].font.size = Pt(10)
                    p.paragraph_format.left_indent = Inches(0.25)

            # 关键技能/技术栈（工作相关）
            if work.get("tech_stack"):
                tech_stack = work.get("tech_stack")
                if isinstance(tech_stack, list):
                    tech_str = ", ".join(tech_stack[:6])  # 最多显示6个
                else:
                    tech_str = str(tech_stack)
                p = doc.add_paragraph()
                p.add_run("技术栈: ").bold = True
                p.add_run(tech_str)
                p.runs[0].font.size = Pt(9)
                p.runs[1].font.size = Pt(9)
                p.runs[1].font.color.rgb = RGBColor(0x64, 0x72, 0x89)

            # 设置段落格式防止孤立行
            for para in doc.paragraphs[-10:]:
                para.paragraph_format.widow_control = True
                para.paragraph_format.keep_with_next = False

            doc.add_paragraph()  # 空行分隔各工作经历

    def _add_work_section_table(self, doc: Document, work_experience: list):
        """使用表格格式添加工作经历"""
        if not work_experience:
            return

        # 添加标题
        p = doc.add_paragraph()
        run = p.add_run("工作经历")
        run.bold = True
        run.font.size = Pt(14)

        # 创建表格
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # 表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "公司"
        hdr_cells[1].text = "职位"
        hdr_cells[2].text = "时间"

        for cell in hdr_cells:
            cell.background_color = RGBColor(0x2B, 0x2B, 0x2B)

        # 添加数据行
        for work in work_experience:
            row_cells = table.add_row().cells
            row_cells[0].text = work.get("company", "")
            row_cells[1].text = work.get("position", "")
            row_cells[2].text = f"{work.get('start_date', '')} - {work.get('end_date', '至今')}"

        doc.add_paragraph()
    
    def _add_project_section(self, doc: Document, projects: list):
        """添加项目经历"""
        p = doc.add_paragraph()
        run = p.add_run("项目经历")
        run.bold = True
        run.font.size = Pt(14)
        
        for project in projects:
            p = doc.add_paragraph()
            p.add_run(project.get('name', '')).bold = True
            if project.get("role"):
                p.add_run(f" ({project['role']})")
            
            if project.get("description"):
                doc.add_paragraph(project["description"])
            
            if project.get("tech_stack"):
                doc.add_paragraph(f"技术栈: {', '.join(project['tech_stack'])}")
        
        doc.add_paragraph()
    
    def _add_skills_section(self, doc: Document, skills: list):
        """添加技能 - 改进版按类别分组"""
        p = doc.add_paragraph()
        run = p.add_run("专业技能")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        p.space_after = Pt(12)

        # 按类别分组技能
        skill_groups = {}
        for skill in skills:
            category = skill.get("category", "其他")
            skill_name = skill.get("name", "")
            if skill_name:
                if category not in skill_groups:
                    skill_groups[category] = []
                skill_groups[category].append(skill_name)

        # 如果有分类，按分类显示
        if skill_groups and len(skill_groups) > 1:
            for category, skill_list in skill_groups.items():
                p = doc.add_paragraph()
                run = p.add_run(f"{category}:")
                run.bold = True
                run.font.size = Pt(10)
                p.add_run(f" {', '.join(skill_list)}")
                p.runs[1].font.size = Pt(10)
        else:
            # 没有分类，全部显示
            skill_names = [skill.get("name", "") for skill in skills if skill.get("name")]
            if skill_names:
                p = doc.add_paragraph()
                # 每行最多5个技能
                chunk_size = 5
                for i in range(0, len(skill_names), chunk_size):
                    chunk = skill_names[i:i + chunk_size]
                    if i == 0:
                        p.add_run(" | ".join(chunk))
                    else:
                        doc.add_paragraph(" | ".join(chunk))

        doc.add_paragraph()
    
    def _add_certifications_section(self, doc: Document, certifications: list):
        """添加证书"""
        p = doc.add_paragraph()
        run = p.add_run("证书荣誉")
        run.bold = True
        run.font.size = Pt(14)
        
        for cert in certifications:
            doc.add_paragraph(f"• {cert.get('name', '')}")
        
        doc.add_paragraph()
